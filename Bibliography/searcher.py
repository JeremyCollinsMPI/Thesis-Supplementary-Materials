import subprocess
from selenium import webdriver
import time
from selenium.webdriver.common.by import By
import dateutil.parser
import ast
import unicodedata
from selenium.webdriver.common.action_chains import ActionChains
from docx import Document
from lxml import etree

from docx import Document
from docx.shared import Inches

import zipfile


file = open('/Users/jeremycollins/Documents/Thesis/Code/Bibliography/raw.txt','r').readlines()
driver=webdriver.Chrome("./chromedriver")

def unique(list):
	result = []
	for x in list:
		if not x in result:
			result.append(x)
	return result
def lookUp(string):
	searchString = string.replace(' ','+')
	driver.get("https://scholar.google.com/scholar?hl=en&as_sdt=0%2C5&q="+string)
	x = driver.find_element_by_class_name('gs_or_cit')
	x.click()
	time.sleep(2)
	y = driver.find_element_by_class_name('gs_citr')
	return y.text

def returnReferences(file):
	results = []	
	for m in xrange(len(file)):
		line = file[m]
		results.append(lookUp(line))
		time.sleep(10)
	return results

def makeIntoBibliography(references):
	references = sorted(references)
	references = unique(references)
	return '\n'.join(references)

def google(string):
	searchString = string.replace(' ','+')
	driver.get("https://www.google.com/search?q="+searchString+'&hl=en'+'&lr=lang_en')
	text = driver.page_source
	textList = text.split('class="r"')
	textList = textList[1:len(textList)]
	textList = [x for x in textList if len(x) > 3]
	results = []
	for i in xrange(2):
		result = textList[i].split("</a>")[0]
		result = result.split('>')[-1]
		print result
		results.append(result)
	print results
	return results

def googleThenLookUp(string):
	results = google(string)
	results2 = []
	for result in results:
		try:
			x = lookUp(result)
			results2.append(x)
		except:
			pass
	return results2
	
def googleTwiceThenLookUp(string):
	results = google(string)
	string2 = ' '.join(results)
	print string2
	print "=----------|"
	results = google(string2)
	results2 = []
	for result in results:
		x = lookUp(result)
		results2.append(x)
	return results2
	
def googleAndLookUpCombined(string):
	results = google(string)
	results2 = []
	for result in results:
		try:
			x = lookUp(result + ' ' +string)
			results2.append(x)
		except:
			pass
	return results2
	 
def returnReferences2(file):
	results = []	
	for m in xrange(len(file)):
		line = file[m]
		results.append(googleThenLookUp(line)[0])
		time.sleep(10)
	return results
def returnReferences3(file):
	results = []	
	for m in xrange(len(file)):
		line = file[m]
		results.append(googleTwiceThenLookUp(line)[0])
	return results
def returnReferences4(file):
	results = []	
	for m in xrange(len(file)):
		line = file[m]
		results.append(googleAndLookUpCombined(line)[0])
	return results




bibliography = open('/Users/jeremycollins/Documents/Thesis/Code/Bibliography/bibliography.txt','w')
bibliography.write(makeIntoBibliography(returnReferences(file)))

# print googleThenLookUp("gray and jordan 2000 austronesian")

	
driver.close()


# def addBibliography():
# 	document = Document('demo.docx')
# 
# 	document.add_heading('Document Title', 0)
# 
# 	def paragraph_replace(self, search, replace):
# 		searchre = re.compile(search)
# 		for paragraph in self.paragraphs:
# 			paragraph_text = paragraph.text
# 			if paragraph_text:
# 				if searchre.search(paragraph_text):
# 					self.clear_paragraph(paragraph)
# 					paragraph.add_run(re.sub(search, replace, paragraph_text))
# 		return paragraph
# 
# 	document.add_paragraph('A plain paragraph having some ')
# 
# 	document.save('demo.docx')
# 
# 	document = Document('demo.docx')
# 
# 
# def get_word_xml(docx_filename):
#    with open(docx_filename) as f:
#       zip = zipfile.ZipFile(f)
#       xml_content = zip.read('word/document.xml')
#       xml_content = xml_content.replace("Title","penguin")
#       
#    return xml_content
   
# def replaceCitations():

# print get_word_xml("/Users/jeremycollins/Documents/Thesis/Code/Thesis.docx")



# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
# 
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='IntenseQuote')
# 
# document.add_paragraph(
#     'first item in unordered list', style='ListBullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='ListNumber'
# )
# 
# document.add_picture('monty-truth.png', width=Inches(1.25))
# 
# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for item in recordset:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc
# 
# document.add_page_break()

# document.save('demo.docx')


# 
# print results

